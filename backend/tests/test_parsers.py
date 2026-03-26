import pytest
import tempfile
import os
from docx import Document as WordDocument
from app.parsers.docx_parser import DocxParser
from app.parsers.markdown_parser import MarkdownParser
from app.parsers.csv_parser import CsvParser


class TestMarkdownParser:
    """Tests for Markdown parser."""

    def test_supports_md_extension(self):
        parser = MarkdownParser()
        assert parser.supports('.md') is True
        assert parser.supports('.markdown') is True
        assert parser.supports('.txt') is False

    def test_parse_with_headings(self, sample_faq_content):
        parser = MarkdownParser()

        # Create temp file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as f:
            f.write(sample_faq_content)
            temp_path = f.name

        try:
            result = parser.parse(temp_path)

            assert result.file_type == 'markdown'
            assert result.confidence == 1.0
            assert len(result.sections) > 0
            assert 'FAQ: Produktinstallation' in result.raw_text
        finally:
            os.unlink(temp_path)

    def test_parse_extracts_sections(self):
        parser = MarkdownParser()
        content = """# Hauptüberschrift

        Einleitung

        ## Unterüberschrift 1

        Inhalt 1

        ## Unterüberschrift 2

        Inhalt 2
        """

        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as f:
            f.write(content)
            temp_path = f.name

        try:
            result = parser.parse(temp_path)

            # Should have sections for headings
            heading_sections = [s for s in result.sections if s.level > 0]
            assert len(heading_sections) >= 2
        finally:
            os.unlink(temp_path)


class TestCsvParser:
    """Tests for CSV parser."""

    def test_supports_csv_extension(self):
        parser = CsvParser()
        assert parser.supports('.csv') is True
        assert parser.supports('.xlsx') is True
        assert parser.supports('.pdf') is False

    def test_parse_csv_creates_sections(self):
        parser = CsvParser()
        content = """name,value,description
        Item1,100,First item
        Item2,200,Second item
        """

        with tempfile.NamedTemporaryFile(mode='w', suffix='.csv', delete=False) as f:
            f.write(content)
            temp_path = f.name

        try:
            result = parser.parse(temp_path)

            assert result.file_type == 'csv'
            assert result.confidence == 1.0
            # Each row becomes a section
            assert len(result.sections) >= 2
            assert result.metadata['row_count'] == 2
        finally:
            os.unlink(temp_path)


class TestDocxParser:
    """Tests for DOCX parser heuristics."""

    def test_parse_detects_heading_like_sections_without_word_styles(self):
        parser = DocxParser()
        doc = WordDocument()
        doc.add_paragraph("Konzept Vertriebsschulung Entmanteler")
        doc.add_paragraph("Das Entmanteler-Prinzip")
        doc.add_paragraph(
            "Dieses Prinzip beschreibt die saubere Vorbereitung verschiedener Kabeltypen "
            "und erklaert die wichtigsten Verkaufsargumente fuer das Vertriebsteam."
        )
        doc.add_paragraph("JOKARI XL")
        doc.add_paragraph(
            "Der JOKARI XL eignet sich fuer groessere Durchmesser und wird in der Schulung "
            "als eigene Produkteinheit mit Nutzenargumentation vorgestellt."
        )

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
            temp_path = handle.name

        try:
            doc.save(temp_path)
            result = parser.parse(temp_path)

            section_titles = [section.title for section in result.sections if section.title]
            assert "Das Entmanteler-Prinzip" in section_titles
            assert "JOKARI XL" in section_titles

            principle_section = next(section for section in result.sections if section.title == "Das Entmanteler-Prinzip")
            assert principle_section.content.startswith("Das Entmanteler-Prinzip")
            assert "Verkaufsargumente" in principle_section.content
        finally:
            os.unlink(temp_path)
